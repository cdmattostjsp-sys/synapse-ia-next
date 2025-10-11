# -*- coding: utf-8 -*-
"""
synapse_chat_vNext.py
POC Synapse.IA – App Streamlit
-------------------------------------------------------------------------------
Melhorias
- Mantém layout atual (branding bar e seções) sem alterar estilos já aprovados.
- Adiciona opção de download do rascunho orientado em .DOCX além do texto.
- Integração com knowledge_base/ via validator_engine_vNext.
- Botão único “Executar Agente” preservado; agora exibe também botões
  “Baixar .DOCX” e “Baixar .MD” após a execução.

Dependências:
  pip install streamlit python-docx

Variáveis de ambiente (no Streamlit Secrets ou env):
  OPENAI_API_KEY   -> chave
  OPENAI_MODEL     -> opcional (ex: gpt-4o-mini)

Arquivos:
  - validator_engine_vNext.py (mesma pasta)
  - knowledge_base/ (na raiz do repo)
-------------------------------------------------------------------------------
"""

from __future__ import annotations
import os
import io
import base64
from datetime import datetime

import streamlit as st
from docx import Document
from docx.shared import Pt
from openai import OpenAI

# engine
from validator_engine_vNext import validate_document

# ----------------------------------------------------------------------------
# Config & helpers
# ----------------------------------------------------------------------------

APP_TITLE = "Synapse.IA"
APP_SUBTITLE = "Tribunal de Justiça de São Paulo"

def _load_api_client():
    api_key = os.getenv("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY", None)
    if not api_key:
        st.error("OPENAI_API_KEY não configurada. Defina em Secrets ou variável de ambiente.")
        st.stop()
    return OpenAI(api_key=api_key)

def branding_bar():
    # Mantém a faixa superior “branding bar” aprovada
    st.markdown(
        """
        <div style="
            background: linear-gradient(180deg,#111418, #0d1014);
            border-bottom: 1px solid #C9A227;
            padding: 18px 20px;">
          <div style="display:flex;align-items:center;gap:16px;">
            <div style="width:42px;height:42px;border-radius:8px;display:flex;align-items:center;justify-content:center;background:#0B2D32;border:1px solid #2A7C86;">
              <span style="font-size:22px;line-height:22px;">🧠</span>
            </div>
            <div>
              <div style="font-size:28px;font-weight:700;color:#ECF2FF;">Synapse.IA</div>
              <div style="font-size:13px;color:#D9C074;margin-top:-2px;">Tribunal de Justiça de São Paulo</div>
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True
    )

def _make_docx_from_markdown(title: str, md_text: str) -> bytes:
    """
    Conversão simples MD -> DOCX (mantém títulos e listas básicas).
    Foco: entregar rascunho utilizável rapidamente no Word.
    """
    doc = Document()
    # Fonte base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_text.splitlines()
    for ln in lines:
        t = ln.strip()
        if t.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(t[2:].strip())
            run.bold = True
            run.font.size = Pt(16)
        elif t.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(t[3:].strip())
            run.bold = True
            run.font.size = Pt(13)
        elif t.startswith("---"):
            doc.add_paragraph().add_run("—" * 60)
        elif t.startswith("> "):
            p = doc.add_paragraph()
            run = p.add_run(t[2:].strip())
            run.italic = True
        elif t.startswith("• ") or t.startswith("- "):
            doc.add_paragraph(t[2:].strip(), style=None).style = doc.styles["List Paragraph"]
        elif t.startswith("<<<INSERIR:"):
            # destaque leve para os marcadores
            p = doc.add_paragraph()
            run = p.add_run(t)
            run.bold = True
        else:
            doc.add_paragraph(t)

    # bytes
    b = io.BytesIO()
    doc.save(b)
    return b.getvalue()

def _download_button_bytes(bytes_data: bytes, filename: str, label: str):
    st.download_button(
        label=label,
        data=bytes_data,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def _download_button_text(text_data: str, filename: str, label: str, mime="text/markdown"):
    st.download_button(
        label=label,
        data=text_data.encode("utf-8", errors="ignore"),
        file_name=filename,
        mime=mime
    )

# ----------------------------------------------------------------------------
# Página
# ----------------------------------------------------------------------------

st.set_page_config(page_title=APP_TITLE, page_icon="🧠", layout="wide")
branding_bar()

st.markdown("")

st.header("🧾 Insumos Manuais")
st.caption("Descreva o objeto, justificativa, requisitos, prazos, critérios etc. Cole abaixo o texto do seu documento (ETP/DFD/TR).")
user_text = st.text_area("", height=220, placeholder="Cole aqui o conteúdo do documento para análise...")

st.header("📤 Upload de Documento (opcional)")
st.caption("Envie PDF, DOCX ou TXT (serão ignorados no processamento atual se o campo acima estiver preenchido).")
upload = st.file_uploader("Drag and drop file here", type=["pdf", "docx", "txt"])

st.header("🤖 Selecionar Agente")
agent = st.selectbox("Escolha o agente:", ["ETP", "DFD", "TR"])

exec_semantic = st.checkbox("Executar validação semântica", value=True)
run = st.button("Executar Agente", type="primary")

if run:
    if not user_text and not upload:
        st.warning("Informe o texto manualmente ou envie um arquivo.")
        st.stop()

    # Prioriza o texto digitado
    raw_text = user_text.strip()
    if not raw_text and upload is not None:
        # leitura simples do upload
        ext = (upload.name.split(".")[-1] or "").lower()
        if ext == "txt":
            raw_text = upload.read().decode("utf-8", errors="ignore")
        else:
            raw_text = upload.read().decode("latin-1", errors="ignore")

    with st.status("Executando agente. Aguarde alguns instantes…", expanded=False):
        try:
            client = _load_api_client()
            result = validate_document(raw_text, agent, client)
        except Exception as e:
            st.error(f"Falha ao executar a validação: {e}")
            st.stop()

    st.success(f"Agente {agent} executado com sucesso!")

    # KPIs
    c1, c2 = st.columns(2)
    with c1:
        st.metric("Score Rígido", f"{result['rigid_score']:.1f}%")
    with c2:
        st.metric("Score Semântico", f"{result['semantic_score']:.1f}%")

    # Tabelas (render simples mantendo padrão atual)
    st.subheader("🧩 Itens Avaliados (Rígidos)")
    # adaptador: mostrar colunas principais
    rigid_rows = []
    for i, it in enumerate(result.get("rigid_result", [])):
        rigid_rows.append({
            "Critério": it.get("descricao", ""),
            "Obrigatório": "✅" if it.get("obrigatorio") else "—",
            "Presente": "✅" if it.get("presente") else "❌"
        })
    st.table(rigid_rows)

    st.subheader("💡 Itens Avaliados (Semânticos)")
    sem_rows = []
    for i, it in enumerate(result.get("semantic_result", [])):
        sem_rows.append({
            "Critério": it.get("descricao", ""),
            "Presente": "✅" if it.get("presente") else "❌",
            "Nota": it.get("adequacao_nota", 0),
            "Justificativa": it.get("justificativa", "")
        })
    st.table(sem_rows)

    # Rascunho Orientado (Markdown)
    st.subheader("📝 Documento orientado")
    st.markdown(result.get("guided_markdown", ""))

    # Downloads
    st.divider()
    st.subheader("⬇️ Exportar rascunho")
    now = datetime.now().strftime("%Y%m%d_%H%M")
    base_name = f"{result.get('guided_doc_title','Rascunho')}_{agent}_{now}"

    # .docx
    docx_bytes = _make_docx_from_markdown(result.get("guided_doc_title", "Rascunho"), result.get("guided_markdown", ""))
    _download_button_bytes(docx_bytes, f"{base_name}.docx", "Baixar .DOCX")

    # .md
    _download_button_text(result.get("guided_markdown", ""), f"{base_name}.md", "Baixar .MD")

    # debug opcional
    with st.expander("🔎 Debug (informações de execução)"):
        st.json(result.get("debug", {}))
